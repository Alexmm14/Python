import requests
import json
from docx import Document


def get_first_contentful_paint(url, api_key):
    # Construir la URL de la API de PageSpeed Insights
    api_url = f"https://www.googleapis.com/pagespeedonline/v5/runPagespeed?url={url}&key={api_key}"

    # Realizar la solicitud a la API
    response = requests.get(api_url)
    data = response.json()

    # Obtener el valor de FIRST_CONTENTFUL_PAINT_MS
    try:
        fcp_value = data["loadingExperience"]["metrics"]["FIRST_CONTENTFUL_PAINT_MS"]["percentile"]
        fcp_value_s = fcp_value / 1000

        return str(fcp_value_s) + ' Segundos'
    except KeyError:
        print("------------------")
        print("La clave FIRST_CONTENTFUL_PAINT_MS no se encuentra en la respuesta.")
        print("------------------")
        return None

"""def print_first_contentful_paint(url, api_key):
    fcp_value = get_first_contentful_paint(url, api_key)

    if fcp_value is not None:
        print("********************")
        print(f"First Contentful Paint (FCP): {fcp_value} segundos")
        print("********************")"""

def get_technology_info_formatted(url, api_key):
    api_url = 'https://whatcms.org/API/Tech'
    params = {'key': api_key, 'url': url}

    response = requests.get(api_url, params=params)
    tech_info_raw = response.text

    try:

        technologysData = []
        technologys = {}
        flag = 1
        hostname_json = json.loads(tech_info_raw)
        for result in hostname_json["results"]:
            
            nombre = result["name"]
            version = result["version"] if "version" in result else "N/A"
            categorias = result["categories"] if "categories" in result else ["N/A"]
            """print("********************")
            print(f"Nombre: {nombre}")
            print(f"Version: {version}")
            print("Categorias:")"""
            technologysData.append(nombre)
            if version:
                technologysData.append(version)
            else:
                version = 'No se pudo obtener la version'
                technologysData.append(version)
            
            categories= []
            for categoria in categorias:
                """print(f" - {categoria}")
                print("-" * 20)"""
                categories.append(categoria)
            technologysData.append(categories)
            technologys['technology' + str(flag)] = technologysData
            technologysData = []
            flag += 1
            
            
        return technologys
    

    except json.JSONDecodeError as e:
        print(f"Error al decodificar JSON: {e}")
        print("\nRespuesta JSON (en bruto):")
        print("------------------")
        print(tech_info_raw)



def getHostsThisSite(url, api_key):
    api_url = 'https://www.who-hosts-this.com/API/Host'
    params = {'key': api_key, 'url': url}

    response = requests.get(api_url, params=params)
    tech_info_raw = response.text

    try:
        hostname_json = json.loads(tech_info_raw)

        for result in hostname_json.get("results", []):
            nombre = result.get("isp_name", "N/A")
            """print("********************")
            print(f"Nombre: {nombre}")
            print("********************")"""
            return nombre

    except json.JSONDecodeError as e:
        print(f"Error al decodificar JSON: {e}")
        print("\nRespuesta JSON (en bruto):")
        print("------------------")
        print(tech_info_raw)





def crear_informe(diccionario):
    # Crear un nuevo documento de Word
    doc = Document()

    # Añadir secciones al documento
    doc.add_heading('Informe de Herramientas y Tiempo de Carga de Página Web', level=0)

    # Introducción
    doc.add_heading('1. Introducción:', level=2)
    doc.add_paragraph('Breve descripción de la página web analizada.')

    # Herramientas Utilizadas
    doc.add_heading('2. Herramientas Utilizadas:', level=2)

    for clave, valores in diccionario.items():
        doc.add_heading(f"{clave}:", level=3)
        
        if clave in ['host', 'tiempo']:
            # Concatenar los valores y agregarlos como un solo párrafo
            doc.add_paragraph(" ".join(map(str, valores)))

        else:
            for valor in valores:
                if isinstance(valor, list):
                    for item in valor:
                        doc.add_paragraph(f"- {item}")
                else:
                    doc.add_paragraph(str(valor))

    # Guardar el documento
    doc.save('Informe_Herramientas_Tiempo_Carga.docx')
    

if __name__ == "__main__":
    # URL que deseas analizar
    url = input("Ingresa la URL que deseas analizar: ")

    # Clave de API de PageSpeed Insights
    api_key_psi = "AIzaSyAL67F5ZAG1eCWeoG5hAOwf2CYO_EfumAE"

    # Clave de API de WhatCMS
    api_key_whatcms = "wq0ec3q7zwkhhko60d7d83856wy0ssvn1lnrbk7w01s327xtjn85up986vja0s70tjrd2l"

    # Obtener información de tecnología a través de la API en formato JSON formateado
    datos = get_technology_info_formatted(url, api_key_whatcms)

    time = get_first_contentful_paint(url, api_key_psi)

    host = getHostsThisSite(url, api_key_whatcms)


    datos['host'] = host

    datos['tiempo'] = str(time)

    print(datos)

    crear_informe(datos)
